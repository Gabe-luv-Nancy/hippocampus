from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# 创建文档
doc = Document()

# 设置中文字体
def set_chinese_font(run, font_name='宋体', font_size=12):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# 文档标题
title = doc.add_heading('国内专精特新企业开发工作计划', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 一、专精特新企业概述
h1 = doc.add_heading('一、专精特新企业概述', level=1)

# 1.1 定义与内涵
h2 = doc.add_heading('1.1 定义与内涵', level=2)
p = doc.add_paragraph()
run = p.add_run('"专精特新"企业是指具有"专业化、精细化、特色化、新颖化"特征的中小企业。其中，国家级专精特新"小巨人"企业是工信部认定的专注于细分市场、创新能力强、市场占有率高、掌握关键核心技术、质量效益优的排头兵企业。')
set_chinese_font(run)

# 1.2 认定层级
h2 = doc.add_heading('1.2 认定层级', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
headers = ['层级', '认定主体', '数量规模']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run, font_size=11, font_name='黑体')

data = [
    ['创新型中小企业', '省级/市級工信部门', '数十万家'],
    ['专精特新中小企业', '省级工信部门', '约10万家'],
    ['专精特新"小巨人"企业', '工信部', '约1.46万家'],
    ['制造业单项冠军', '工信部', '约1200家']
]
for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx+1].cells[col_idx].text = cell_data
        for paragraph in table.rows[row_idx+1].cells[col_idx].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, font_size=11)

# 二、总量统计与区域分布
h1 = doc.add_heading('二、总量统计与区域分布', level=1)

# 2.1 国家级专精特新"小巨人"企业数量
h2 = doc.add_heading('2.1 国家级专精特新"小巨人"企业数量（截至2024年底）', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
headers = ['统计口径', '数量', '数据来源']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run, font_size=11, font_name='黑体')

data = [
    ['国家级"小巨人"', '1.46万家', '工信部'],
    ['省级专精特新中小企业', '约10万家', '各省工信厅'],
    ['创新型中小企业', '数十万家', '各市工信局']
]
for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx+1].cells[col_idx].text = cell_data
        for paragraph in table.rows[row_idx+1].cells[col_idx].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, font_size=11)

# 2.2 区域分布
h2 = doc.add_heading('2.2 区域分布（前十城市）', level=2)
table = doc.add_table(rows=11, cols=4)
table.style = 'Table Grid'
headers = ['排名', '城市', '"小巨人"数量', '特点']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run, font_size=11, font_name='黑体')

data = [
    ['1', '深圳', '1333家', '全国第一，"专精特新第一城"'],
    ['2', '北京', '约1000家', '科技创新中心'],
    ['3', '上海', '约800家', '金融+科技融合'],
    ['4', '苏州', '约700家', '制造业强市'],
    ['5', '宁波', '约600家', '制造业单项冠军'],
    ['6', '杭州', '约500家', '数字经济'],
    ['7', '青岛', '约450家', '家电产业链'],
    ['8', '成都', '约400家', '西部中心'],
    ['9', '武汉', '约350家', '光谷科技'],
    ['10', '合肥', '约300家', '新兴产业']
]
for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx+1].cells[col_idx].text = cell_data
        for paragraph in table.rows[row_idx+1].cells[col_idx].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, font_size=11)

# 2.3 行业分布
h2 = doc.add_heading('2.3 行业分布', level=2)
p = doc.add_paragraph()
run = p.add_run('"小巨人"企业主要集中在以下战略性新兴产业：')
set_chinese_font(run)

industries = [
    ('新一代信息技术', '占比约28%'),
    ('高端装备制造', '占比约22%'),
    ('新材料', '占比约18%'),
    ('新能源', '占比约12%'),
    ('生物医药', '占比约10%'),
    ('节能环保', '占比约10%')
]
table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
for i, header in enumerate(['行业', '占比']):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run, font_size=11, font_name='黑体')

for row_idx, (industry, ratio) in enumerate(industries):
    table.rows[row_idx+1].cells[0].text = industry
    table.rows[row_idx+1].cells[1].text = ratio
    for col_idx in range(2):
        for paragraph in table.rows[row_idx+1].cells[col_idx].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, font_size=11)

# 三、开发工作目标
h1 = doc.add_heading('三、开发工作目标', level=1)

# 3.1 总体目标
h2 = doc.add_heading('3.1 总体目标', level=2)
p = doc.add_paragraph()
run = p.add_run('通过3-5年努力，构建覆盖全国主要经济区域的专精特新企业培育网络，形成"发现-培育-认定-服务-壮大"的完整工作体系。')
set_chinese_font(run)

# 3.2 量化指标
h2 = doc.add_heading('3.2 量化指标', level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
for i, header in enumerate(['指标', '2025年目标', '2027年目标']):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run, font_size=11, font_name='黑体')

data = [
    ['走访企业数量', '500家', '2000家'],
    ['纳入培育库企业', '200家', '800家'],
    ['省级专精特新认定', '50家', '200家'],
    ['国家级"小巨人"认定', '10家', '40家'],
    ['制造业单项冠军', '2家', '8家']
]
for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx+1].cells[col_idx].text = cell_data
        for paragraph in table.rows[row_idx+1].cells[col_idx].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, font_size=11)

# 四、开发工作路径
h1 = doc.add_heading('四、开发工作路径', level=1)

# 4.1 企业筛查与发现
h2 = doc.add_heading('4.1 企业筛查与发现', level=2)
p = doc.add_paragraph()
run = p.add_run('数据来源渠道：')
set_chinese_font(run, font_name='黑体')

channels = [
    '政府部门：工信局、科技局、发改委名录',
    '园区管理：经开区、高新区、特色产业园入驻企业',
    '行业协会：各细分领域协会会员单位',
    '金融机构：银行保荐的优质客户',
    '公开数据：上市辅导企业、招投标数据、专利数据库'
]
for channel in channels:
    p = doc.add_paragraph(channel)
    for run in p.runs:
        set_chinese_font(run)

p = doc.add_paragraph()
run = p.add_run('筛选标准：')
set_chinese_font(run, font_name='黑体')

standards = [
    '营业收入：2000万-4亿元',
    '从业人员：300人以下',
    '细分市场：市占率排名前三',
    '研发投入：营收占比≥5%',
    '知识产权：发明专利≥5项'
]
for standard in standards:
    p = doc.add_paragraph(standard)
    for run in p.runs:
        set_chinese_font(run)

# 4.2 分层分类培育
h2 = doc.add_heading('4.2 分层分类培育', level=2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
for i, header in enumerate(['层级', '培育周期', '重点服务']):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run, font_size=11, font_name='黑体')

data = [
    ['创新型中小企业', '1-2年', '诊断对标、政策宣导'],
    ['专精特新中小企业', '1-2年', '融资对接、产学研'],
    ['"小巨人"企业', '2-3年', '上市辅导、产业链协同']
]
for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx+1].cells[col_idx].text = cell_data
        for paragraph in table.rows[row_idx+1].cells[col_idx].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, font_size=11)

# 4.3 核心开发动作
h2 = doc.add_heading('4.3 核心开发动作', level=2)

h3 = doc.add_heading('（一）政府渠道对接', level=3)
items = [
    '与各地工信局建立合作关系',
    '参与中小企业服务月活动',
    '承接政府采购项目供应商开发'
]
for item in items:
    p = doc.add_paragraph(item)
    for run in p.runs:
        set_chinese_font(run)

h3 = doc.add_heading('（二）园区深度覆盖', level=3)
items = [
    '重点园区派驻服务专员',
    '举办专精特新政策宣讲会',
    '建立园区企业培育库'
]
for item in items:
    p = doc.add_paragraph(item)
    for run in p.runs:
        set_chinese_font(run)

h3 = doc.add_heading('（三）产业链垂直开发', level=3)
items = [
    '围绕链主企业（央企、国企、上市公司）梳理上下游',
    '挖掘供应链关键环节的专精特新企业',
    '推动产业链上下游协同创新'
]
for item in items:
    p = doc.add_paragraph(item)
    for run in p.runs:
        set_chinese_font(run)

h3 = doc.add_heading('（四）数字化工具应用', level=3)
items = [
    '企业大数据平台筛查',
    'AI舆情监控发现潜力企业',
    'CRM系统跟踪培育进度'
]
for item in items:
    p = doc.add_paragraph(item)
    for run in p.runs:
        set_chinese_font(run)

# 五、服务体系构建
h1 = doc.add_heading('五、服务体系构建', level=1)

h2 = doc.add_heading('5.1 政策服务', level=2)
services = [
    '专精特新认定政策解读',
    '研发费用加计扣除申报',
    '科技创新券申请',
    '人才引进政策匹配'
]
for service in services:
    p = doc.add_paragraph(service)
    for run in p.runs:
        set_chinese_font(run)

h2 = doc.add_heading('5.2 融资服务', level=2)
services = [
    '银行贷款（专精特新贷）',
    '股权投资对接',
    '供应链金融',
    '上市辅导（北交所为主）'
]
for service in services:
    p = doc.add_paragraph(service)
    for run in p.runs:
        set_chinese_font(run)

h2 = doc.add_heading('5.3 产学研对接', level=2)
services = [
    '高校技术成果转化',
    '实验室共享',
    '专家顾问聘请',
    '关键技术联合攻关'
]
for service in services:
    p = doc.add_paragraph(service)
    for run in p.runs:
        set_chinese_font(run)

h2 = doc.add_heading('5.4 市场拓展', level=2)
services = [
    '产业链供需对接会',
    '政府/国企采购对接',
    '头部企业供应链准入',
    '国内外展会组织'
]
for service in services:
    p = doc.add_paragraph(service)
    for run in p.runs:
        set_chinese_font(run)

# 六、考核激励机制
h1 = doc.add_heading('六、考核激励机制', level=1)

# 6.1 开发团队考核
h2 = doc.add_heading('6.1 开发团队考核', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
for i, header in enumerate(['指标', '权重', '考核周期']):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_chinese_font(run, font_size=11, font_name='黑体')

data = [
    ['新增培育企业数量', '30%', '季度'],
    ['省级认定通过率', '25%', '半年度'],
    ['国家级认定通过率', '25%', '年度'],
    ['客户满意度', '20%', '季度']
]
for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx+1].cells[col_idx].text = cell_data
        for paragraph in table.rows[row_idx+1].cells[col_idx].paragraphs:
            for run in paragraph.runs:
                set_chinese_font(run, font_size=11)

# 6.2 激励政策
h2 = doc.add_heading('6.2 激励政策', level=2)
p = doc.add_paragraph()
run = p.add_run('认定成功奖励：')
set_chinese_font(run, font_name='黑体')
run = p.add_run('国家级"小巨人"企业服务团队奖励5万元/家')
set_chinese_font(run)

p = doc.add_paragraph()
run = p.add_run('年度超额完成奖：')
set_chinese_font(run, font_name='黑体')
run = p.add_run('超额完成目标的团队给予额外奖金')
set_chinese_font(run)

p = doc.add_paragraph()
run = p.add_run('优秀案例分享：')
set_chinese_font(run, font_name='黑体')
run = p.add_run('最佳开发案例给予荣誉表彰')
set_chinese_font(run)

# 七、2025年度工作计划
h1 = doc.add_heading('七、2025年度工作计划', level=1)

# 7.1-7.4
quarters = [
    ('7.1 Q1：基础建设期', [
        '搭建企业数据库（目标：纳入1000家企业信息）',
        '组建专业服务团队（目标：5人以上）',
        '签订3个重点园区合作协议',
        '完成政策产品培训'
    ]),
    ('7.2 Q2：市场开拓期', [
        '举办5场专精特新政策宣讲会',
        '走访目标企业200家',
        '纳入培育库企业50家',
        '完成省级认定申报15家'
    ]),
    ('7.3 Q3：成果转化期', [
        '省级专精特新认定通过20家',
        '国家级"小巨人"申报5家',
        '举办产业链对接活动2场',
        '融资服务落地企业10家'
    ]),
    ('7.4 Q4：总结提升期', [
        '国家级"小巨人"认定通过3家',
        '完成年度目标复盘',
        '制定下年度工作计划',
        '团队能力提升培训'
    ])
]

for title, tasks in quarters:
    h2 = doc.add_heading(title, level=2)
    for task in tasks:
        p = doc.add_paragraph(task)
        for run in p.runs:
            set_chinese_font(run)

# 八、风险控制
h1 = doc.add_heading('八、风险控制', level=1)

# 8.1 主要风险
h2 = doc.add_heading('8.1 主要风险', level=2)
risks = [
    '政策变化风险：认定标准调整',
    '竞争加剧风险：同行机构争夺客户',
    '服务能力风险：专业人才流失',
    '企业经营风险：培育企业出现经营问题'
]
for risk in risks:
    p = doc.add_paragraph(risk)
    for run in p.runs:
        set_chinese_font(run)

# 8.2 应对措施
h2 = doc.add_heading('8.2 应对措施', level=2)
measures = [
    '建立政策跟踪机制，及时调整工作策略',
    '加强客户关系维护，提升服务粘性',
    '完善团队激励机制，核心人才梯队建设',
    '强化尽职调查，控制企业准入风险'
]
for measure in measures:
    p = doc.add_paragraph(measure)
    for run in p.runs:
        set_chinese_font(run)

# 脚注
p = doc.add_paragraph()
run = p.add_run('*本计划基于2024-2025年公开数据编制，具体数据以官方最新公布为准。*')
set_chinese_font(run)
run.font.italic = True

# 保存文档到Windows桌面
output_path = r'C:\Users\GabetopZ\Desktop\专精特新企业开发工作计划.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
